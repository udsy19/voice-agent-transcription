"""Export history and logs in multiple formats: JSON, TXT, Word, PDF."""

import json
import os
import time
import shutil
import tempfile
from pathlib import Path
from logger import get_logger

log = get_logger("exporter")


def export_json(history: list[dict], path: str) -> str:
    """Export history as formatted JSON."""
    with open(path, "w") as f:
        json.dump(history, f, indent=2, ensure_ascii=False)
    return path


def export_txt(history: list[dict], path: str) -> str:
    """Export history as plain text."""
    with open(path, "w") as f:
        f.write(f"Muse — History Export\n")
        f.write(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"{'=' * 50}\n\n")
        for i, entry in enumerate(history, 1):
            f.write(f"#{i}  [{entry.get('ts', '')}]  {entry.get('app', 'Unknown')}\n")
            f.write(f"Raw:     {entry.get('raw', '')}\n")
            cleaned = entry.get('cleaned', '')
            if cleaned and cleaned != entry.get('raw', ''):
                f.write(f"Cleaned: {cleaned}\n")
            dur = entry.get('duration', 0)
            if dur:
                f.write(f"Duration: {dur}s\n")
            f.write(f"{'-' * 40}\n\n")
    return path


def export_word(history: list[dict], path: str) -> str:
    """Export history as a Word document (.docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.error("python-docx not installed — run: pip install python-docx")
        raise ImportError("python-docx required for Word export. Install: pip install python-docx")

    doc = Document()

    # Title
    title = doc.add_heading("Muse — History", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generated {time.strftime('%B %d, %Y at %I:%M %p')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(135, 114, 99)

    doc.add_paragraph()  # spacer

    for i, entry in enumerate(history, 1):
        # Entry header
        h = doc.add_heading(level=2)
        h_run = h.add_run(f"#{i}")
        h_run.font.color.rgb = RGBColor(28, 14, 4)

        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"{entry.get('ts', '')}  •  {entry.get('app', 'Unknown')}  •  {entry.get('duration', 0)}s")
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(135, 114, 99)

        # Raw text
        if entry.get('raw'):
            raw_p = doc.add_paragraph()
            raw_label = raw_p.add_run("Raw: ")
            raw_label.font.size = Pt(9)
            raw_label.font.color.rgb = RGBColor(135, 114, 99)
            raw_text = raw_p.add_run(entry['raw'])
            raw_text.font.size = Pt(10)
            raw_text.font.color.rgb = RGBColor(100, 100, 100)

        # Cleaned text (main content)
        cleaned = entry.get('cleaned', '')
        if cleaned:
            p = doc.add_paragraph()
            run = p.add_run(cleaned)
            run.font.size = Pt(12)

        doc.add_paragraph()  # spacer between entries

    doc.save(path)
    return path


def export_pdf(history: list[dict], path: str) -> str:
    """Export history as PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    except ImportError:
        log.error("reportlab not installed — run: pip install reportlab")
        raise ImportError("reportlab required for PDF export. Install: pip install reportlab")

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        'VATitle', parent=styles['Title'],
        fontSize=22, textColor=HexColor('#1C0E04'), spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        'VAMeta', parent=styles['Normal'],
        fontSize=9, textColor=HexColor('#877263'), spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        'VARaw', parent=styles['Normal'],
        fontSize=9, textColor=HexColor('#999999'), spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        'VACleaned', parent=styles['Normal'],
        fontSize=11, textColor=HexColor('#1C0E04'), spaceAfter=8,
        leading=16,
    ))
    styles.add(ParagraphStyle(
        'VAEntry', parent=styles['Heading2'],
        fontSize=13, textColor=HexColor('#1C0E04'), spaceAfter=4,
    ))

    story = []

    story.append(Paragraph("Muse — History", styles['VATitle']))
    story.append(Paragraph(
        f"Generated {time.strftime('%B %d, %Y at %I:%M %p')}",
        styles['VAMeta']
    ))
    story.append(Spacer(1, 12))

    for i, entry in enumerate(history, 1):
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=HexColor('#ECE3D8'), spaceAfter=8))

        story.append(Paragraph(f"#{i}", styles['VAEntry']))
        story.append(Paragraph(
            f"{entry.get('ts', '')} &bull; {entry.get('app', 'Unknown')} &bull; {entry.get('duration', 0)}s",
            styles['VAMeta']
        ))

        if entry.get('raw'):
            # Escape XML special chars for reportlab
            raw = entry['raw'].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(f"Raw: {raw}", styles['VARaw']))

        cleaned = entry.get('cleaned', '')
        if cleaned:
            cleaned = cleaned.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            cleaned = cleaned.replace('\n', '<br/>')
            story.append(Paragraph(cleaned, styles['VACleaned']))

        story.append(Spacer(1, 6))

    doc.build(story)
    return path


def export_logs(log_dir: str, output_path: str) -> str:
    """Zip all log files for bug reporting."""
    if os.path.exists(log_dir):
        shutil.make_archive(output_path.replace('.zip', ''), 'zip', log_dir)
        return output_path
    raise FileNotFoundError(f"Log directory not found: {log_dir}")
